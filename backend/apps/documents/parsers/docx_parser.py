"""DOCX document parser."""

from docx import Document as DocxDocument

from apps.documents.parsers.base import BaseParser, ParsedDocument


class DocxParser(BaseParser):
    SUPPORTED_TYPES = {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }

    def supports(self, mime_type: str) -> bool:
        return mime_type in self.SUPPORTED_TYPES

    def parse(self, file_path: str) -> ParsedDocument:
        doc = DocxDocument(file_path)
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract tables
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                paragraphs.append("\n".join(rows))

        full_text = "\n\n".join(paragraphs)
        metadata = {
            "parser": "docx",
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
        }
        return ParsedDocument(text=full_text, metadata=metadata, page_count=1)
